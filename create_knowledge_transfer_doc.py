from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_document():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Second Brain")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(8, 16, 40)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(0)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Knowledge Transfer")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(126, 137, 172)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(6)

    # Tagline
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run("Capture expert knowledge. Generate training automatically.")
    tagline_run.italic = True
    tagline_run.font.size = Pt(10)
    tagline_run.font.color.rgb = RGBColor(100, 100, 100)
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.space_after = Pt(12)

    # What is Second Brain?
    heading1 = doc.add_paragraph()
    h1_run = heading1.add_run("What is Second Brain?")
    h1_run.bold = True
    h1_run.font.size = Pt(12)
    h1_run.font.color.rgb = RGBColor(8, 16, 40)
    heading1.space_after = Pt(4)

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Second Brain is a knowledge transfer tool that captures expertise from your organization's documents, "
        "identifies knowledge gaps, collects expert answers, and automatically generates training videos. "
        "It transforms scattered information into a searchable, shareable knowledge base."
    )
    intro_run.font.size = Pt(10)
    intro.space_after = Pt(10)

    # Step 1: Connect Your Sources
    step1 = doc.add_paragraph()
    s1_run = step1.add_run("Step 1: Connect Your Sources")
    s1_run.bold = True
    s1_run.font.size = Pt(11)
    s1_run.font.color.rgb = RGBColor(8, 16, 40)
    step1.space_after = Pt(4)

    # Integration table
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Category"
    header_cells[1].text = "Integrations"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, "FFE2BF")

    # Data rows
    data = [
        ("Communication", "Slack, Gmail"),
        ("Cloud Storage", "Box, GitHub"),
        ("Office Tools", "Microsoft PowerPoint, Microsoft Excel"),
        ("Research", "PubMed, ResearchGate, Google Scholar"),
        ("Manual", "Drag & drop file uploads"),
        ("Custom", "Custom integrations for your lab or organization"),
    ]

    for i, (category, integrations) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = category
        row[1].text = integrations
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph().space_after = Pt(6)

    # Step 2: Automatic Analysis
    step2 = doc.add_paragraph()
    s2_run = step2.add_run("Step 2: Automatic Analysis")
    s2_run.bold = True
    s2_run.font.size = Pt(11)
    s2_run.font.color.rgb = RGBColor(8, 16, 40)
    step2.space_after = Pt(2)

    step2_content = doc.add_paragraph()
    step2_content.add_run("Once connected, documents sync automatically. The AI parses all content, extracts key information, "
        "and identifies knowledge gaps—missing rationale, undefined processes, unclear definitions, "
        "evidence gaps, and more. No manual work required.").font.size = Pt(10)
    step2_content.space_after = Pt(8)

    # Step 3: Fill Knowledge Gaps
    step3 = doc.add_paragraph()
    s3_run = step3.add_run("Step 3: Fill Knowledge Gaps")
    s3_run.bold = True
    s3_run.font.size = Pt(11)
    s3_run.font.color.rgb = RGBColor(8, 16, 40)
    step3.space_after = Pt(2)

    step3_content = doc.add_paragraph()
    step3_content.add_run("The system generates up to 30 prioritized questions based on detected gaps. "
        "Subject matter experts can answer via text or voice recording (automatically transcribed). "
        "Each answer is immediately embedded into the knowledge base and becomes searchable.").font.size = Pt(10)
    step3_content.space_after = Pt(8)

    # Step 4: Generate Training Videos
    step4 = doc.add_paragraph()
    s4_run = step4.add_run("Step 4: Generate Training Videos")
    s4_run.bold = True
    s4_run.font.size = Pt(11)
    s4_run.font.color.rgb = RGBColor(8, 16, 40)
    step4.space_after = Pt(2)

    step4_content = doc.add_paragraph()
    step4_content.add_run("Once knowledge gaps are filled, the system automatically creates training materials. "
        "The Gamma API generates professional presentations, and Azure Text-to-Speech adds voiceovers. "
        "The result: HD training videos (1920x1080) ready to download and share with your team.").font.size = Pt(10)
    step4_content.space_after = Pt(8)

    # Step 5: Use Your Knowledge Base
    step5 = doc.add_paragraph()
    s5_run = step5.add_run("Step 5: Use & Share Your Knowledge")
    s5_run.bold = True
    s5_run.font.size = Pt(11)
    s5_run.font.color.rgb = RGBColor(8, 16, 40)
    step5.space_after = Pt(2)

    step5_content = doc.add_paragraph()
    step5_content.add_run("Use the chat interface to ask questions and get AI-powered answers sourced from your documents "
        "and expert input. Share training videos across your organization. "
        "The knowledge base grows smarter with every answered gap.").font.size = Pt(10)
    step5_content.space_after = Pt(8)

    # The Human Touch
    human_heading = doc.add_paragraph()
    human_run = human_heading.add_run("The Human Touch")
    human_run.bold = True
    human_run.font.size = Pt(11)
    human_run.font.color.rgb = RGBColor(8, 16, 40)
    human_heading.space_after = Pt(2)

    human_content = doc.add_paragraph()
    human_content.add_run("Second Brain enhances—not replaces—human expertise. AI identifies gaps and generates questions, "
        "but the answers come from your subject matter experts. The system captures institutional knowledge "
        "that only humans possess: context, nuance, and real-world experience. "
        "Think of it as a tool that amplifies your experts, not one that replaces them.").font.size = Pt(10)
    human_content.space_after = Pt(8)

    # Security
    security_heading = doc.add_paragraph()
    security_run = security_heading.add_run("Security & Privacy")
    security_run.bold = True
    security_run.font.size = Pt(11)
    security_run.font.color.rgb = RGBColor(8, 16, 40)
    security_heading.space_after = Pt(2)

    security_content = doc.add_paragraph()
    security_content.add_run("Your data is protected with enterprise-grade security: multi-tenant isolation ensures "
        "each organization's data is completely separate; OAuth 2.0 authentication secures all integrations; "
        "JWT tokens protect every API request; and all data is encrypted in transit and at rest. "
        "Your knowledge stays yours.").font.size = Pt(10)
    security_content.space_after = Pt(10)

    # Visual Flow
    flow = doc.add_paragraph()
    flow_run = flow.add_run("Knowledge Transfer Flow")
    flow_run.bold = True
    flow_run.font.size = Pt(10)
    flow_run.font.color.rgb = RGBColor(8, 16, 40)
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow.space_after = Pt(4)

    flow_diagram = doc.add_paragraph()
    flow_text = flow_diagram.add_run("Connect Sources  →  Auto-Analyze  →  Answer Gaps  →  Generate Training  →  Share & Use")
    flow_text.font.size = Pt(10)
    flow_text.bold = True
    flow_text.font.color.rgb = RGBColor(8, 16, 40)
    flow_diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc.save('/Users/rishitjain/Downloads/2nd-brain/Knowledge_Transfer.docx')
    print("Document created: Knowledge_Transfer.docx")

if __name__ == "__main__":
    create_document()
